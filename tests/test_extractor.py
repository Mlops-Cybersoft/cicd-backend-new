from io import BytesIO

from docx import Document

from app.services.extractor import extract_text


def test_extract_utf8_text() -> None:
    pages = extract_text("Quy định nghỉ phép".encode(), "policy.txt")
    assert pages[0].text == "Quy định nghỉ phép"
    assert pages[0].page_number is None


def test_extract_docx_paragraphs() -> None:
    document = Document()
    document.add_paragraph("Điều 1. Phạm vi áp dụng")
    document.add_paragraph("Điều 2. Trách nhiệm thi hành")
    buffer = BytesIO()
    document.save(buffer)

    pages = extract_text(buffer.getvalue(), "regulation.docx")
    assert "Điều 1" in pages[0].text
    assert "Điều 2" in pages[0].text
